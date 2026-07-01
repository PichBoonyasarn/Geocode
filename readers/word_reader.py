import hashlib
import io
from docx import Document
from PIL import Image


def extract_text_and_images(docx_path: str, max_pages: int = 2) -> tuple[str, list[bytes]]:
    """
    Extract text and map-candidate images from a Word document.
    Word files have no real page concept; we approximate 2 pages as the first 100 paragraphs.
    Images are pulled from all embedded relationships (they're not page-scoped in python-docx).
    """
    doc = Document(docx_path)

    # Text: first ~100 paragraphs (approx. 2 pages)
    para_texts = [p.text for p in doc.paragraphs[:100] if p.text.strip()]

    # Also grab cell text from tables on the first page
    table_texts: list[str] = []
    for table in doc.tables[:10]:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_texts.append(cell.text.strip())

    text = "\n".join(para_texts + table_texts)

    # Images: extract from part relationships
    images: list[bytes] = []
    seen_hashes: set[str] = set()

    for rel in doc.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            img_bytes = rel.target_part.blob
        except Exception:
            continue

        img_hash = hashlib.md5(img_bytes).hexdigest()
        if img_hash in seen_hashes:
            continue
        seen_hashes.add(img_hash)

        try:
            pil_img = Image.open(io.BytesIO(img_bytes))
            w, h = pil_img.size
            if w >= 200 and h >= 200:
                images.append(img_bytes)
        except Exception:
            continue

    return text, images
